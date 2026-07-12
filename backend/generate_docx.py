import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    # Helper to set table cell background color
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def create_documentation_docx(output_path="TransitOps_Platform_Documentation.docx"):
    doc = docx.Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # Styles definition
    styles = doc.styles
    
    # Title style
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("TransitOps Platform")
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(26)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(15, 23, 42) # Slate 900
    
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle_p.add_run("In-Depth Operational Reference & Workflow Matrix Guide")
    sub_run.font.name = 'Arial'
    sub_run.font.size = Pt(12)
    sub_run.font.italic = True
    sub_run.font.color.rgb = RGBColor(100, 116, 139) # Slate 500
    
    doc.add_paragraph() # Spacer
    
    # Document content
    def add_h1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = RGBColor(15, 23, 42) # Slate 900
        
    def add_h2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(37, 99, 235) # Electric Blue
        
    def add_body(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(51, 65, 85) # Slate 700
        return p

    def add_bullet(text):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(51, 65, 85)
        
    # --- 1. Executive Summary ---
    add_h1("1. Executive Summary")
    add_body(
        "TransitOps is a smart, full-stack transport operations platform designed to coordinate fleet management, "
        "trip logistics, driver dispatch, asset maintenance, and operational logs. By enforcing strict verification "
        "rules and embedding Gemini AI context views, the system provides real-time synchronization between "
        "dispatchers, drivers, maintenance crews, and administrators."
    )
    
    add_body(
        "The system has been built on a modular decoupled architecture where a React Vite Tailwind CSS frontend "
        "interacts with a Python Django REST Framework backend API using JSON Web Tokens (JWT) for authentication."
    )
    
    # --- 2. RBAC Matrix ---
    add_h1("2. User Access & Role Matrix")
    add_body(
        "The platform operates on a Role-Based Access Control (RBAC) hierarchy. Menu structures, page actions, "
        "and API endpoints are dynamically filtered based on these credentials:"
    )
    
    # Roles Table
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Role Type'
    hdr_cells[1].text = 'UI View Focus'
    hdr_cells[2].text = 'Permissions & Scope'
    
    # Make header bold and slate background
    for cell in hdr_cells:
        set_cell_background(cell, "0F172A")
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.font.size = Pt(9.5)
                
    role_rows = [
        ("Administrator (ADMIN)", "Control Center Theme (Violet)\nFull nav view + Admin Panel links", "• Platform-wide database access\n• User management & driver profile verification\n• Analytics exports & cost auditing"),
        ("Dispatcher (DISPATCHER)", "Dispatch Console Theme (Blue)\nTrips, Drivers, and Vehicles registries", "• Trip registry creation & dispatch assignment\n• Payload limit validations\n• Driver eligibility & license expiry overrides"),
        ("Maintenance Manager (MAINTENANCE)", "Fleet Workshop Theme (Amber)\nMaintenance requests, servicing records", "• Repairs logging & completion hooks\n• Vehicle condition state changes\n• OOS (Out-of-service) triggers"),
        ("Fleet Driver (DRIVER)", "Driver Portal Theme (Emerald)\nPersonal dispatch logs, fuel & expense logs", "• Route GPS & status updates\n• Fuel receipts & strictly increasing odometer logs\n• Trip expense logs submission")
    ]
    
    for idx, row in enumerate(role_rows):
        cells = table.rows[idx + 1].cells
        cells[0].text = row[0]
        cells[1].text = row[1]
        cells[2].text = row[2]
        
        # Color even rows grey
        bg_color = "F8FAFC" if idx % 2 == 1 else "FFFFFF"
        for col_idx, cell in enumerate(cells):
            set_cell_background(cell, bg_color)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
                    r.font.color.rgb = RGBColor(51, 65, 85)
                    if col_idx == 0:
                        r.font.bold = True
                        r.font.color.rgb = RGBColor(15, 23, 42)
                        
    doc.add_paragraph() # spacer
    
    # --- 3. Operational Workflows ---
    add_h1("3. Core Operational Workflows")
    
    add_h2("3.1 Trip Dispatch Lifecycle")
    add_body("The primary workflow path for logistics operations follows a state-machine progression:")
    add_bullet("Scheduling: The Dispatcher matches an available vehicle and driver for a specific route. Behind the scenes, validations confirm the driver's license is active (not expired) and that the cargo weight does not exceed the vehicle's max payload capacity.")
    add_bullet("Activation: Upon starting the trip, the state shifts to In Progress. This locks the vehicle status to On Trip and the driver to On Trip.")
    add_bullet("Resolution: The driver updates route logs. When complete, the dispatcher updates the state to Completed (which prompts for actual trip mileage update) or Cancelled. The vehicles and drivers are automatically returned to the Available roster.")
    
    add_h2("3.2 Asset Maintenance & Safety Lifecycle")
    add_body("Ensures all vehicles are operating safely through automatic condition locking:")
    add_bullet("Service Request: Maintenance managers trigger a repair log for a specific vehicle.")
    add_bullet("Locking: The vehicle status is automatically updated to Maintenance (or Out of Service if critical) to block dispatchers from scheduling it for any new trips.")
    add_bullet("Release: Once maintenance works are completed and logged, the vehicle status reverts back to Available automatically.")
    
    add_h2("3.3 Financial Logging & Verification Loop")
    add_body("Keeps fuel and general expenses synchronized with strict math checks:")
    add_bullet("Fuel Log Odometer Check: Odometer inputs must exceed the vehicle's last recorded odometer reading. Successful logging automatically updates the vehicle's odometer value.")
    add_bullet("Cost Aggregation: Expense logs (tolls, lodging, etc.) are matched with dispatch IDs, rolling up into the monthly financial audit views.")
    
    add_h2("3.4 Live Location Tracking & External Navigation")
    add_body("Integrates real-world routing and GPS tracking interfaces dynamically:")
    add_bullet("Light-Mode Vector Mapping: Utilizes Leaflet.js and CartoDB Voyager light tiles to draw road paths by geocoding origin/destination parameters using Nominatim and calculating OSRM coordinates.")
    add_bullet("Speed & Progress Simulation: Active dispatches show real-time animated indicator progress, variable speed control, and continuous ETA estimations.")
    add_bullet("Google Maps External Link: Provides deep-linked 'Navigate (Google Maps)' controls to load origin/destination address variables directly in external routing tabs.")
    
    add_h2("3.5 Dedicated Live Tracking Dashboard")
    add_body("Provides a dedicated space in the main navigation for fleet-wide real-time tracking:")
    add_bullet("Roster Panel: Lists all active dispatches currently in-progress with vehicle and driver assignments.")
    add_bullet("Map Center: Clicking on any active dispatch dynamically focuses the light-mode map on its route path with live telemetry simulation.")

    # --- 4. AI Assistant ---
    add_h1("4. AI Copilot Integration")
    add_body(
        "The platform embeds a Gemini 2.5 Flash assistant. When a user requests operations analysis, "
        "the system dynamically constructs a real-time fleet snapshot (counts of active trips, available drivers, "
        "maintenance tasks, and vehicles) and injects it as a system prompt instruction. This guarantees that "
        "answers regarding fleet utilization and scheduling are grounded in the actual database state."
    )
    add_body(
        "The output UI includes full Markdown rendering, enabling structured layouts, bold highlights, list items, "
        "and clean code block displays natively within the Copilot chat drawer."
    )
    
    # Save document
    doc.save(output_path)
    print("DOCX Generation complete.")

if __name__ == "__main__":
    create_documentation_docx("TransitOps_Platform_Documentation.docx")
